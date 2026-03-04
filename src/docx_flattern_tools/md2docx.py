"""
Markdown 到 Docx 转换器
"""

import os
import re
import subprocess
import tempfile
from pathlib import Path
from docx import Document
from docx.shared import Inches


def parse_markdown_content(lines):
    """解析整个 markdown 内容，识别各种元素包括 mermaid 代码块"""
    elements = []
    i = 0

    while i < len(lines):
        line = lines[i].rstrip("\n")

        # 检查是否为 mermaid 代码块开始
        if line.strip() == "```mermaid":
            # 找到代码块结束
            code_lines = []
            i += 1
            while i < len(lines) and lines[i].strip() != "```":
                code_lines.append(lines[i].rstrip("\n"))
                i += 1

            # 跳过结束的 ```
            if i < len(lines):
                i += 1

            elements.append({"type": "mermaid", "content": "\n".join(code_lines)})
            continue

        # 检查是否为其他代码块（跳过）
        elif line.strip().startswith("```"):
            # 找到代码块结束
            i += 1
            while i < len(lines) and lines[i].strip() != "```":
                i += 1
            # 跳过结束的 ```
            if i < len(lines):
                i += 1
            continue

        # 检查是否为标题
        heading_match = re.match(r"^(#{1,6})\s+(.*)", line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2)
            elements.append({"type": "heading", "level": level, "text": text})

        # 检查是否为空行
        elif line.strip() == "":
            elements.append({"type": "empty"})

        # 处理普通段落
        else:
            elements.append({"type": "paragraph", "text": line})

        i += 1

    return elements


def check_mmdc_available():
    """检查系统中是否可用 mmdc 命令"""
    try:
        result = subprocess.run(
            ["mmdc", "--version"], capture_output=True, text=True, timeout=10
        )
        return result.returncode == 0
    except (FileNotFoundError, OSError, subprocess.TimeoutExpired):
        return False


def render_mermaid_to_image(mermaid_code, output_path):
    """
    使用 mermaid-cli 将 Mermaid 代码渲染为 PNG 图片

    :param mermaid_code: Mermaid 代码字符串
    :param output_path: 输出图片路径
    :return: 是否成功渲染
    """
    # 检查 mmdc 是否可用
    if not check_mmdc_available():
        print("警告: 未找到 mermaid-cli (mmdc)。将回退到显示源码。")
        print("如需渲染 Mermaid 图表，请安装: npm install -g @mermaid-js/mermaid-cli")
        return False

    try:
        # 创建临时文件存储 mermaid 代码
        with tempfile.NamedTemporaryFile(
            mode="w", suffix=".mmd", delete=False
        ) as temp_file:
            temp_file.write(mermaid_code)
            temp_file_path = temp_file.name

        # 使用 mmdc 命令渲染图片
        result = subprocess.run(
            [
                "mmdc",
                "-i",
                temp_file_path,
                "-o",
                output_path,
                "--backgroundColor",
                "white",
            ],
            capture_output=True,
            text=True,
            timeout=30,
        )

        # 清理临时文件
        os.unlink(temp_file_path)

        if result.returncode == 0:
            return True
        else:
            print(f"Mermaid 渲染失败: {result.stderr}")
            return False

    except (
        FileNotFoundError,
        OSError,
        subprocess.SubprocessError,
        subprocess.TimeoutExpired,
    ) as e:
        print(f"渲染 Mermaid 图表时出错: {e}")
        return False


def create_docx_from_md(md_path: str, docx_path: str):
    """
    将 markdown 文件转换为 docx 格式。

    :param md_path: 输入 markdown 文件的路径
    :param docx_path: 输出 docx 文件的路径
    """
    # 创建新的文档
    doc = Document()

    # 读取 markdown 文件
    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.readlines()

    # 解析内容
    elements = parse_markdown_content(lines)

    # 创建临时目录存储 mermaid 图片
    with tempfile.TemporaryDirectory() as temp_dir:
        image_counter = 0

        for element in elements:
            if element["type"] == "heading":
                # 添加标题
                heading = doc.add_heading(element["text"], level=element["level"])

            elif element["type"] == "paragraph":
                # 添加段落
                paragraph = doc.add_paragraph(element["text"])

            elif element["type"] == "empty":
                # 添加空行
                doc.add_paragraph()

            elif element["type"] == "mermaid":
                # 渲染 mermaid 为图片并添加到文档
                image_counter += 1
                image_path = os.path.join(temp_dir, f"mermaid_{image_counter}.png")

                if render_mermaid_to_image(element["content"], image_path):
                    if os.path.exists(image_path):
                        # 添加图片到文档
                        doc.add_paragraph("Mermaid 图表:")
                        doc.add_picture(image_path, width=Inches(6))
                        doc.add_paragraph()  # 添加空行
                    else:
                        # 如果图片渲染失败，回退到代码显示
                        paragraph = doc.add_paragraph()
                        run = paragraph.add_run("Mermaid 图表 (渲染失败，显示源码):\n")
                        run.bold = True
                        code_run = paragraph.add_run(element["content"])
                        code_run.font.name = "Courier New"
                else:
                    # 如果渲染失败，回退到代码显示
                    paragraph = doc.add_paragraph()
                    run = paragraph.add_run("Mermaid 图表 (渲染失败，显示源码):\n")
                    run.bold = True
                    code_run = paragraph.add_run(element["content"])
                    code_run.font.name = "Courier New"

        # 保存文档
        doc.save(docx_path)
        print(f"已将 {md_path} 转换为 {docx_path}")


def main():
    """命令行使用的主要函数。"""
    import argparse

    parser = argparse.ArgumentParser(description="将 markdown 文件转换为 docx")
    parser.add_argument("input", help="输入 markdown 文件路径")
    parser.add_argument("-o", "--output", help="输出 docx 文件路径")

    args = parser.parse_args()

    # 确定输出文件路径
    if not args.output:
        output_path = args.input.replace(".md", ".docx")
    else:
        output_path = args.output

    # 如果输出目录不存在，则创建它
    output_dir = os.path.dirname(output_path)
    if output_dir and not os.path.exists(output_dir):
        os.makedirs(output_dir)

    # 执行转换
    create_docx_from_md(args.input, output_path)


if __name__ == "__main__":
    main()
