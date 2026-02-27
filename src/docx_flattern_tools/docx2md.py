"""
Docx 到 Markdown 转换器
"""
import os
from docx import Document


def convert_paragraph(paragraph):
    """处理段落文本，保留基本格式"""
    para_text = ""
    for run in paragraph.runs:
        text = run.text
        if run.bold:
            text = f"**{text}**"
        if run.italic:
            text = f"*{text}*"
        if run.underline:
            text = f"<u>{text}</u>"
        para_text += text
    return para_text


def convert_table_to_md(table):
    """将表格转换为 markdown 格式"""
    md_table = []

    # 处理表头
    header_row = table.rows[0]
    header_cells = [cell.text.strip() for cell in header_row.cells]
    md_table.append("| " + " | ".join(header_cells) + " |")
    md_table.append("| " + " | ".join(["---"] * len(header_cells)) + " |")

    # 处理数据行
    for row in table.rows[1:]:
        row_data = [cell.text.strip().replace("\n", "<br>") for cell in row.cells]
        md_table.append("| " + " | ".join(row_data) + " |")

    return "\n".join(md_table)


def convert_docx_to_md(docx_path: str, md_path: str):
    """
    将 docx 文件转换为 markdown 格式。

    :param docx_path: 输入 docx 文件的路径
    :param md_path: 输出 markdown 文件的路径
    """
    # 加载 docx 文档
    doc = Document(docx_path)

    # 准备 markdown 内容
    md_content = []

    # 遍历文档中的所有段落和表格
    # 按照在文档中出现的顺序处理元素
    for element in doc.element.body:
        if element.tag.endswith('p'):  # 段落
            # 查找对应的段落对象
            paragraph = None
            for p in doc.paragraphs:
                if p._element is element:
                    paragraph = p
                    break

            if paragraph is not None:
                para_text = convert_paragraph(paragraph)

                # 检查段落样式是否为标题
                if paragraph.style.name.startswith('Heading'):
                    level = int(paragraph.style.name.split()[-1])
                    md_content.append(f"{'#' * level} {para_text}")
                else:
                    md_content.append(para_text)

        elif element.tag.endswith('tbl'):  # 表格
            # 查找对应的表格对象
            table = None
            for t in doc.tables:
                if t._element is element:
                    table = t
                    break

            if table is not None:
                md_content.append(convert_table_to_md(table))

        # 在每个元素后添加空行
        md_content.append("")

    # 将 markdown 内容写入文件
    with open(md_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(md_content))

    print(f"已将 {docx_path} 转换为 {md_path}")


def main():
    """命令行使用的主要函数。"""
    import argparse

    parser = argparse.ArgumentParser(description="将 docx 文件转换为 markdown")
    parser.add_argument("input", help="输入 docx 文件路径")
    parser.add_argument("-o", "--output", help="输出 markdown 文件路径")

    args = parser.parse_args()

    # 确定输出文件路径
    if not args.output:
        output_path = args.input.replace('.docx', '.md')
    else:
        output_path = args.output

    # 如果输出目录不存在，则创建它
    output_dir = os.path.dirname(output_path)
    if output_dir and not os.path.exists(output_dir):
        os.makedirs(output_dir)

    # 执行转换
    convert_docx_to_md(args.input, output_path)


if __name__ == "__main__":
    main()